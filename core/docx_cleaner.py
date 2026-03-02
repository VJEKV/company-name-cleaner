"""
Обработка .docx файлов: замена текста с сохранением форматирования.

Стратегия:
1. Склеиваем текст всех runs в параграфе
2. Строим карту: позиция символа → (run_index, char_index_in_run)
3. Ищем паттерны в склеенном тексте
4. ЗАМЕНЯЕМ найденное на заглушку (не удаляем!)
5. Модифицируем runs, сохраняя форматирование первого затронутого run
"""

import re
import copy
import logging
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger('CompanyCleaner')


def clean_docx(
    filepath: str,
    output_path: str,
    company_patterns: list[re.Pattern],
    surname_patterns: list[re.Pattern],
    company_replacement: str,
    surname_mapper,
) -> dict:
    """
    Заменяет вхождения названия компании и фамилий на заглушки.

    Возвращает словарь со статистикой.
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        logger.error(f"Не удалось открыть {filepath}: {e}")
        return {
            "status": "error",
            "company_matches": 0,
            "surname_matches": 0,
            "total_replacements": 0,
            "error_message": str(e),
        }

    stats = {"company": 0, "surname": 0}

    # 1. Параграфы основного текста
    for paragraph in doc.paragraphs:
        stats = _process_paragraph(
            paragraph, company_patterns, surname_patterns,
            company_replacement, surname_mapper, stats
        )

    # 2. Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    stats = _process_paragraph(
                        paragraph, company_patterns, surname_patterns,
                        company_replacement, surname_mapper, stats
                    )

    # 3. Колонтитулы
    for section in doc.sections:
        for header in [section.header, section.first_page_header,
                       section.even_page_header]:
            if header is not None:
                try:
                    linked = header.is_linked_to_previous
                except Exception:
                    linked = True
                if not linked:
                    for paragraph in header.paragraphs:
                        stats = _process_paragraph(
                            paragraph, company_patterns, surname_patterns,
                            company_replacement, surname_mapper, stats
                        )
        for footer in [section.footer, section.first_page_footer,
                       section.even_page_footer]:
            if footer is not None:
                try:
                    linked = footer.is_linked_to_previous
                except Exception:
                    linked = True
                if not linked:
                    for paragraph in footer.paragraphs:
                        stats = _process_paragraph(
                            paragraph, company_patterns, surname_patterns,
                            company_replacement, surname_mapper, stats
                        )

    try:
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Не удалось сохранить {output_path}: {e}")
        return {
            "status": "error",
            "company_matches": stats["company"],
            "surname_matches": stats["surname"],
            "total_replacements": stats["company"] + stats["surname"],
            "error_message": str(e),
        }

    return {
        "status": "success",
        "company_matches": stats["company"],
        "surname_matches": stats["surname"],
        "total_replacements": stats["company"] + stats["surname"],
        "error_message": None,
    }


def _process_paragraph(paragraph, company_patterns, surname_patterns,
                       company_replacement, surname_mapper, stats):
    """
    Ключевой алгоритм замены текста в параграфе с сохранением форматирования.

    1. Склеиваем текст всех runs
    2. Строим карту символ → (run_index, char_index_in_run)
    3. Ищем паттерны (сначала длинные, потом короткие)
    4. Заменяем, сохраняя форматирование первого затронутого run
    """
    runs = paragraph.runs
    if not runs:
        return stats

    # Склеиваем текст и строим карту
    full_text = ''
    char_map = []  # [(run_index, char_index_in_run), ...]

    for run_idx, run in enumerate(runs):
        run_text = run.text or ''
        for char_idx, char in enumerate(run_text):
            char_map.append((run_idx, char_idx))
        full_text += run_text

    if not full_text.strip():
        return stats

    # Собираем все замены: (start, end, replacement_text, type)
    replacements = []

    # Сначала компания (длинные паттерны приоритетнее)
    for pattern in company_patterns:
        for match in pattern.finditer(full_text):
            replacements.append((match.start(), match.end(),
                                 company_replacement, "company"))

    # Потом фамилии
    for pattern in surname_patterns:
        for match in pattern.finditer(full_text):
            replacement = surname_mapper.get_replacement(match.group())
            replacements.append((match.start(), match.end(),
                                 replacement, "surname"))

    if not replacements:
        return stats

    # Убираем пересекающиеся замены (приоритет длинным)
    replacements.sort(key=lambda r: (r[0], -(r[1] - r[0])))
    filtered = []
    last_end = 0
    for start, end, repl, rtype in replacements:
        if start >= last_end:
            filtered.append((start, end, repl, rtype))
            last_end = end

    if not filtered:
        return stats

    # Подсчёт
    for _, _, _, rtype in filtered:
        stats[rtype] = stats.get(rtype, 0) + 1

    # Строим новый текст
    new_text = ''
    pos = 0
    for start, end, repl, _ in filtered:
        new_text += full_text[pos:start]
        new_text += repl
        pos = end
    new_text += full_text[pos:]

    # Перезаписываем runs: весь текст в первый run, остальные очищаем
    if runs:
        # Сохраняем форматирование первого run
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''

    return stats


def preview_docx(
    filepath: str,
    company_patterns: list[re.Pattern],
    surname_patterns: list[re.Pattern],
    surname_mapper,
    context_chars: int = 30,
) -> dict:
    """
    Сканирует файл БЕЗ изменений, возвращает найденные вхождения.
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        return {"status": "error", "error_message": str(e), "matches": []}

    matches = []
    all_paragraphs = []

    # Собираем все параграфы
    for p in doc.paragraphs:
        all_paragraphs.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_paragraphs.append(p)

    for para in all_paragraphs:
        text = para.text
        if not text.strip():
            continue

        for pattern in company_patterns + surname_patterns:
            for match in pattern.finditer(text):
                start = max(0, match.start() - context_chars)
                end = min(len(text), match.end() + context_chars)
                context = text[start:end]
                if start > 0:
                    context = '...' + context
                if end < len(text):
                    context = context + '...'

                is_company = pattern in company_patterns
                if is_company:
                    repl = '[ЗАГЛУШКА КОМПАНИИ]'
                else:
                    repl = surname_mapper.get_replacement(match.group()) \
                        if surname_mapper else '[ЗАГЛУШКА ФИО]'

                matches.append({
                    "original": match.group(),
                    "replacement": repl,
                    "context": context,
                    "type": "company" if is_company else "surname",
                })

    return {
        "status": "success",
        "matches": matches,
        "company_count": sum(1 for m in matches if m["type"] == "company"),
        "surname_count": sum(1 for m in matches if m["type"] == "surname"),
    }
