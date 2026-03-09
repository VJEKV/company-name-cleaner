"""
Автоматическое обнаружение персональных данных и реквизитов в тексте.

Категории:
- Фамилии (ФИО, с инициалами, без)
- Организации (ООО, ПАО, ИП и т.д.)
- Города
- Реквизиты (ИНН, ОГРН, КПП, БИК, р/с, к/с)
- Паспортные данные, СНИЛС
- Телефоны, email
- Почтовые адреса
"""

import re
from dataclasses import dataclass, field

from core.cities_db import RUSSIAN_CITIES
from core.whitelist import is_whitelisted_org, is_whitelisted_in_context


# ── Типы сущностей ──────────────────────────────────────────

ENTITY_SURNAME = "surname"
ENTITY_ORGANIZATION = "organization"
ENTITY_CITY = "city"
ENTITY_INN = "inn"
ENTITY_OGRN = "ogrn"
ENTITY_KPP = "kpp"
ENTITY_BIK = "bik"
ENTITY_ACCOUNT = "account"
ENTITY_SNILS = "snils"
ENTITY_PASSPORT = "passport"
ENTITY_PHONE = "phone"
ENTITY_EMAIL = "email"
ENTITY_ADDRESS = "address"


@dataclass
class DetectedEntity:
    """Обнаруженная сущность в тексте."""
    start: int
    end: int
    text: str
    entity_type: str
    replacement: str = ""
    confidence: float = 1.0  # 0..1


# ── Справочники ─────────────────────────────────────────────

# Русские мужские имена (именительный падеж)
MALE_NAMES: frozenset[str] = frozenset({
    "Абрам", "Аврам", "Адам", "Адриан", "Аким", "Александр", "Алексей",
    "Альберт", "Анатолий", "Андрей", "Антон", "Аркадий", "Арсений",
    "Артём", "Артемий", "Артур", "Афанасий",
    "Богдан", "Борис", "Бронислав",
    "Вадим", "Валентин", "Валерий", "Василий", "Вениамин", "Виктор",
    "Виталий", "Владимир", "Владислав", "Всеволод", "Вячеслав",
    "Гавриил", "Геннадий", "Георгий", "Герман", "Глеб", "Григорий",
    "Давид", "Даниил", "Данила", "Демьян", "Денис", "Дмитрий",
    "Евгений", "Егор", "Емельян", "Ефим",
    "Захар", "Зиновий",
    "Иван", "Игнат", "Игорь", "Илларион", "Илья", "Иннокентий",
    "Кирилл", "Климент", "Константин", "Кузьма",
    "Лаврентий", "Лев", "Леонид", "Леонтий", "Лука",
    "Макар", "Максим", "Марат", "Марк", "Матвей", "Мирон", "Михаил",
    "Назар", "Никита", "Николай", "Никон",
    "Олег", "Остап",
    "Павел", "Пётр", "Платон", "Прохор",
    "Радий", "Рафаил", "Ринат", "Роберт", "Родион", "Роман",
    "Ростислав", "Руслан", "Рустам",
    "Савелий", "Святослав", "Семён", "Сергей", "Станислав", "Степан",
    "Тарас", "Тимофей", "Тимур", "Тихон",
    "Ульян", "Устин",
    "Фёдор", "Федот", "Феликс", "Филипп",
    "Харитон", "Христофор",
    "Эдуард", "Эльдар", "Эмиль", "Эрик",
    "Юлиан", "Юрий",
    "Яков", "Ян", "Ярослав",
})

# Русские женские имена
FEMALE_NAMES: frozenset[str] = frozenset({
    "Агата", "Аделина", "Алевтина", "Александра", "Алина", "Алиса",
    "Алла", "Альбина", "Анастасия", "Ангелина", "Анжела", "Анна",
    "Антонина", "Арина",
    "Белла", "Богдана",
    "Валентина", "Валерия", "Варвара", "Василиса", "Вера", "Вероника",
    "Виктория", "Виолетта",
    "Галина", "Глафира",
    "Дарина", "Дарья", "Диана", "Дина",
    "Ева", "Евгения", "Евдокия", "Екатерина", "Елена", "Елизавета",
    "Жанна",
    "Зинаида", "Злата", "Зоя",
    "Инга", "Инна", "Ираида", "Ирина",
    "Карина", "Кира", "Клавдия", "Клара", "Кристина", "Ксения",
    "Лада", "Лариса", "Лидия", "Лилия", "Лина", "Любовь", "Людмила",
    "Мавра", "Маргарита", "Марина", "Мария", "Марта", "Милана", "Мирослава",
    "Надежда", "Наталия", "Наталья", "Нелли", "Нина",
    "Оксана", "Олеся", "Ольга",
    "Пелагея", "Полина", "Прасковья",
    "Рада", "Раиса", "Регина", "Римма", "Роза", "Руслана",
    "Светлана", "Снежана", "Софья", "Стелла",
    "Таисия", "Тамара", "Татьяна",
    "Ульяна",
    "Фаина",
    "Христина",
    "Элеонора", "Элина", "Эльвира", "Эмилия", "Эмма",
    "Юлия", "Юнона",
    "Яна", "Ярослава",
})

ALL_FIRST_NAMES = MALE_NAMES | FEMALE_NAMES

# Окончания фамилий (от длинных к коротким)
SURNAME_ENDINGS_MALE = (
    'ский', 'цкий', 'ской', 'цкой',
    'ов', 'ев', 'ёв', 'ин', 'ын',
    'их', 'ых',
    'ко', 'ук', 'юк', 'як', 'ак', 'ик',
    'ец', 'ич', 'уш',
)

SURNAME_ENDINGS_FEMALE = (
    'ская', 'цкая',
    'ова', 'ева', 'ёва', 'ина', 'ына',
    'их', 'ых',
    'ко', 'ук', 'юк', 'як', 'ак',
    'ец', 'ич',
)

ALL_SURNAME_ENDINGS = tuple(sorted(
    set(SURNAME_ENDINGS_MALE + SURNAME_ENDINGS_FEMALE),
    key=len, reverse=True
))

# Слова-ложные срабатывания (выглядят как фамилии, но не фамилии)
FALSE_POSITIVE_SURNAMES: frozenset[str] = frozenset({
    # Общие слова на -ов/-ева/-ин и т.д.
    "Готов", "Готова", "Основ", "Основа", "Здоров", "Здорова",
    "Нов", "Нова", "Новой", "Стар", "Таков", "Такова",
    "Каков", "Какова", "Сущностей",
    # Месяцы и дни
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
    "Понедельник", "Вторник", "Среда", "Четверг", "Пятница",
    "Суббота", "Воскресенье",
    # Юридические термины
    "Истец", "Ответчик", "Заявитель", "Должник", "Кредитор",
    "Исполнитель", "Заказчик", "Поставщик", "Покупатель",
    "Арендатор", "Арендодатель", "Подрядчик", "Субподрядчик",
    "Субподрядчиков", "Субподрядчика", "Субподрядчику",
    "Работник", "Работодатель", "Продавец",
    "Подрядчиков", "Подрядчика", "Подрядчику",
    # Должности
    "Директор", "Президент", "Генеральный", "Главный", "Начальник",
    "Заместитель", "Руководитель", "Секретарь", "Бухгалтер",
    "Менеджер", "Специалист", "Инженер", "Консультант",
    # Типичные слова документов
    "Приложение", "Дополнение", "Изменение", "Условие",
    "Договор", "Контракт", "Соглашение", "Протокол",
    "Основание", "Положение", "Решение", "Определение",
    "Заключение", "Уведомление", "Требование",
    "Обязательство", "Представитель", "Действующий",
    "Настоящий", "Указанный", "Данный", "Который",
    "Однако", "Поскольку", "Поэтому", "Именно",
    "Также", "Кроме", "После", "Перед", "Между",
    "Нижеподписавшиеся",
    # Прилагательные и причастия (частые ложные срабатывания)
    "Возможных", "Возможный", "Возможная", "Возможное", "Возможного",
    "Возможной", "Возможным", "Возможном", "Возможные", "Возможно",
    "Иных", "Иной", "Иного", "Иному", "Иным", "Ином", "Иные", "Иная",
    "Имеющих", "Имеющий", "Имеющая", "Имеющее", "Имеющего",
    "Имеющей", "Имеющим", "Имеющем", "Имеющие",
    "Состоящих", "Состоящий", "Состоящая", "Состоящее", "Состоящего",
    "Состоящей", "Состоящим", "Состоящем", "Состоящие",
    "Действующих", "Действующий", "Действующая", "Действующее",
    "Действующего", "Действующей", "Действующим", "Действующем",
    "Следующих", "Следующий", "Следующая", "Следующее", "Следующего",
    "Следующей", "Следующим", "Следующем", "Следующие",
    "Предыдущих", "Предыдущий", "Предыдущая", "Предыдущее",
    "Указанных", "Указанный", "Указанная", "Указанное", "Указанного",
    "Надлежащий", "Надлежащая", "Надлежащее", "Надлежащего",
    "Надлежащих", "Надлежащим", "Ненадлежащий",
    "Соответствующих", "Соответствующий", "Соответствующая",
    "Существующих", "Существующий", "Существующая",
    "Настоящих", "Настоящего", "Настоящей", "Настоящим", "Настоящем",
    "Текущих", "Текущий", "Текущая", "Текущее", "Текущего",
    "Вышестоящий", "Вышестоящая", "Вышестоящих",
    "Нижестоящий", "Нижестоящая", "Нижестоящих",
    "Графиком", "График", "Графика", "Графику", "Графике", "Графики",
    "Юридический", "Юридическая", "Юридическое", "Юридического",
    "Юридической", "Юридическим", "Юридическом", "Юридических",
    "Юридические", "Юридически",
    "Физический", "Физическая", "Физическое", "Физического",
    "Физической", "Физическим", "Физическом", "Физических",
    "Фактический", "Фактическая", "Фактическое", "Фактического",
    "Фактической", "Фактическим", "Фактическом", "Фактических",
    "Технический", "Технического", "Технической", "Техническим",
    "Техническом", "Технических", "Технические", "Техническая",
    "Практический", "Практического", "Практических",
    "Экономический", "Экономического", "Экономических",
    "Строительный", "Строительного", "Строительных",
    "Металлургический", "Металлургического", "Металлургических",
    "Народов", "Народный", "Народная", "Народное", "Народного",
    "Народной", "Народным", "Народном", "Народных", "Народные",
    # Географические прилагательные
    "Нижегородский", "Нижегородская", "Нижегородское", "Нижегородского",
    "Нижегородской", "Нижегородским", "Нижегородском", "Нижегородских",
    "Ставропольский", "Ставропольская", "Ставропольское",
    "Ставропольского", "Ставропольской", "Ставропольским",
    "Краснодарский", "Краснодарская", "Краснодарского", "Краснодарской",
    "Московский", "Московская", "Московское", "Московского",
    "Московской", "Московским", "Московском", "Московских",
    "Ленинградский", "Ленинградская", "Ленинградского", "Ленинградской",
    "Свердловский", "Свердловская", "Свердловского", "Свердловской",
    "Новосибирский", "Новосибирская", "Новосибирского",
    "Челябинский", "Челябинская", "Челябинского",
    "Самарский", "Самарская", "Самарского",
    "Волгоградский", "Волгоградская", "Волгоградского",
    "Ростовский", "Ростовская", "Ростовского",
    "Тульский", "Тульская", "Тульского",
    "Рязанский", "Рязанская", "Рязанского",
    "Тверской", "Тверская", "Тверского",
    "Калужский", "Калужская", "Калужского",
    "Брянский", "Брянская", "Брянского",
    "Воронежский", "Воронежская", "Воронежского",
    "Саратовский", "Саратовская", "Саратовского",
    "Пензенский", "Пензенская", "Пензенского",
    "Тамбовский", "Тамбовская", "Тамбовского",
    "Ульяновский", "Ульяновская", "Ульяновского",
    "Оренбургский", "Оренбургская", "Оренбургского",
    "Башкирский", "Башкирская", "Башкирского",
    "Красноярский", "Красноярская", "Красноярского",
    "Хабаровский", "Хабаровская", "Хабаровского",
    "Приморский", "Приморская", "Приморского",
    "Алтайский", "Алтайская", "Алтайского",
    "Забайкальский", "Забайкальская", "Забайкальского",
    "Камчатский", "Камчатская", "Камчатского",
    "Сахалинский", "Сахалинская", "Сахалинского",
    "Мурманский", "Мурманская", "Мурманского",
    "Архангельский", "Архангельская", "Архангельского",
    "Вологодский", "Вологодская", "Вологодского",
    "Ярославский", "Ярославская", "Ярославского",
    "Костромской", "Костромская", "Костромского",
    "Ивановский", "Ивановская", "Ивановского",
    "Владимирский", "Владимирская", "Владимирского",
    "Смоленский", "Смоленская", "Смоленского",
    "Курский", "Курская", "Курского",
    "Белгородский", "Белгородская", "Белгородского",
    "Липецкий", "Липецкая", "Липецкого",
    "Орловский", "Орловская", "Орловского",
    "Кемеровский", "Кемеровская", "Кемеровского",
    "Омский", "Омская", "Омского",
    "Томский", "Томская", "Томского",
    "Тюменский", "Тюменская", "Тюменского",
    "Иркутский", "Иркутская", "Иркутского",
    "Амурский", "Амурская", "Амурского",
    "Пермский", "Пермская", "Пермского",
    "Кировский", "Кировская", "Кировского",
    # Прочие общие слова
    "Итого", "Всего", "Вместе", "Далее", "Выше", "Ниже",
    "Россия", "Российский", "Российская", "Российское",
    "Российской", "Российских", "Российскому", "Российским",
    "Советский", "Советская", "Советское", "Советского",
    "Федерации", "Федеральный", "Федеральная", "Федеральное",
    "Федерального", "Федеральной",
    "Общество", "Компания", "Предприятие", "Учреждение",
    "Организация", "Филиал", "Отделение", "Управление",
    "Департамент", "Министерство",
    "Северный", "Южный", "Западный", "Восточный",
    "Центральный", "Верхний", "Нижний",
    "Большой", "Малый", "Новый", "Старый",
    "Областной", "Областная", "Областное", "Областного",
    "Областной", "Областным", "Областном", "Областных",
    "Краевой", "Краевая", "Краевое", "Краевого",
    "Районный", "Районная", "Районное", "Районного",
    "Городской", "Городская", "Городское", "Городского",
    "Сельский", "Сельская", "Сельское", "Сельского",
    "Муниципальный", "Муниципальная", "Муниципальное",
    "Муниципального", "Муниципальной", "Муниципальных",
    "Автономный", "Автономная", "Автономное", "Автономного",
    "Республиканский", "Республиканская", "Республиканского",
    # Существительные на -ов/-ин/-ский
    "Проектов", "Объектов", "Документов", "Контрактов",
    "Платежей", "Вариантов", "Фактов", "Результатов",
    "Участков", "Районов", "Городов", "Номеров",
    "Товаров", "Продуктов", "Рисков", "Сроков",
    "Расходов", "Доходов", "Налогов", "Взносов",
    "Договоров", "Актов", "Счетов", "Вопросов",
    "Органов", "Законов", "Кодексов", "Стандартов",
    "Условий", "Оснований", "Положений", "Решений",
    "Требований", "Обязательств", "Полномочий",
})

# Орг. формы
ORG_FORMS = (
    'ООО', 'ОАО', 'ЗАО', 'ПАО', 'АО', 'НАО',
    'ИП', 'ФГУП', 'МУП', 'ГУП', 'НКО', 'АНО',
    'ФГБУ', 'ФГКУ', 'ГБУ', 'МБУ', 'МБОУ', 'ГБОУ',
)

ORG_FORMS_FULL = (
    'Общество с ограниченной ответственностью',
    'Открытое акционерное общество',
    'Закрытое акционерное общество',
    'Публичное акционерное общество',
    'Акционерное общество',
    'Индивидуальный предприниматель',
)


# ── Regex-паттерны ───────────────────────────────────────────

# Инициалы: А.Б. или А. Б. или А.Б
_INIT = r'[А-ЯЁ]\.\s?[А-ЯЁ]\.?'

# ФИО полное: Фамилия Имя Отчество
_PATRONYMIC = r'[А-ЯЁ][а-яё]+(?:ович(?:а|у|ем|е)?|евич(?:а|у|ем|е)?|ёвич(?:а|у|ем|е)?|ич(?:а|у|ем|е)?|овн[аыеуой]|евн[аыеуой]|ёвн[аыеуой]|ичн[аыеуой]|иничн[аыеуой])'

RE_FULL_NAME = re.compile(
    rf'([А-ЯЁ][а-яё]{{2,}})\s+([А-ЯЁ][а-яё]+)\s+({_PATRONYMIC})'
)

# Тот же паттерн но для имён с маленькой буквы после фамилии (родительный и т.д.)
# "Козловой Елены Николаевны"
RE_FULL_NAME_LOWER = re.compile(
    rf'([А-ЯЁ][а-яё]{{2,}})\s+([А-ЯЁ][а-яё]+)\s+({_PATRONYMIC})',
    re.IGNORECASE,
)

# Фамилия + инициалы: Иванов А.В. / Иванов А. В.
RE_SURNAME_INITIALS = re.compile(
    rf'([А-ЯЁ][а-яё]{{2,}})\s+({_INIT})'
)

# Инициалы + фамилия: А.В. Иванов
RE_INITIALS_SURNAME = re.compile(
    rf'({_INIT})\s+([А-ЯЁ][а-яё]{{2,}})'
)

# Организация: ООО «Название» или ООО "Название"
_ORG_ALT = '|'.join(re.escape(f) for f in ORG_FORMS)
_QUOTE_OPEN = r'[«"\u201c\u201e\'"]'
_QUOTE_CLOSE = r'[»"\u201d\'"]'

RE_ORG_QUOTED = re.compile(
    rf'(?:{_ORG_ALT})\s*{_QUOTE_OPEN}([^»"\u201d\'"]+){_QUOTE_CLOSE}'
)

# Орг. форма полная + «Название»
_ORG_FULL_ALT = '|'.join(re.escape(f) for f in ORG_FORMS_FULL)
RE_ORG_FULL_QUOTED = re.compile(
    rf'(?:{_ORG_FULL_ALT})\s*{_QUOTE_OPEN}([^»"\u201d\'"]+){_QUOTE_CLOSE}',
    re.IGNORECASE,
)

# ИП Фамилия И.О.
RE_IP = re.compile(
    rf'ИП\s+([А-ЯЁ][а-яё]{{2,}}(?:\s+{_INIT}|\s+[А-ЯЁ][а-яё]+\s+{_PATRONYMIC})?)'
)

# Реквизиты
RE_INN = re.compile(r'ИНН\s*:?\s*(\d{10}(?:\d{2})?)\b')
RE_OGRN = re.compile(r'ОГРН(?:ИП)?\s*:?\s*(\d{13,15})\b')
RE_KPP = re.compile(r'КПП\s*:?\s*(\d{9})\b')
RE_BIK = re.compile(r'БИК\s*:?\s*(\d{9})\b')
RE_ACCOUNT = re.compile(
    r'(?:р/с|р\.с\.|расч[её]тный\s+сч[её]т|к/с|к\.с\.|корр[.]?\s*сч[её]т|'
    r'лицевой\s+сч[её]т|л/с)\s*:?\s*(\d{20})\b',
    re.IGNORECASE,
)
# Просто 20 цифр подряд после "счёт" / "счет"
RE_ACCOUNT_BARE = re.compile(
    r'(?:сч[её]т|№\s*сч[её]та)\s*:?\s*(\d{20})\b',
    re.IGNORECASE,
)

# СНИЛС: 123-456-789 01 или 12345678901
RE_SNILS = re.compile(
    r'(?:СНИЛС|Страховой\s+номер)\s*:?\s*'
    r'(\d{3}[-\s]?\d{3}[-\s]?\d{3}[-\s]?\d{2})\b',
    re.IGNORECASE,
)

# Паспорт: серия 12 34 номер 567890, паспорт 1234 567890
RE_PASSPORT = re.compile(
    r'(?:паспорт|серия)\s*:?\s*'
    r'(\d{2}\s?\d{2})\s*(?:№|номер|н-р)?\s*:?\s*(\d{6})',
    re.IGNORECASE,
)
RE_PASSPORT_ISSUED = re.compile(
    r'(?:выдан|выдано|выд\.)\s*:?\s*(.{10,80}?)(?:\d{2}\.\d{2}\.\d{4})',
    re.IGNORECASE,
)

# Телефон
RE_PHONE = re.compile(
    r'(?:тел\.?|телефон|моб\.?|факс|т\.)\s*:?\s*'
    r'([+]?[78][\s(-]*\d{3}[\s)-]*\d{3}[\s-]*\d{2}[\s-]*\d{2})',
    re.IGNORECASE,
)
RE_PHONE_BARE = re.compile(
    r'(?<!\d)([+]7[\s(-]*\d{3}[\s)-]*\d{3}[\s-]*\d{2}[\s-]*\d{2})(?!\d)'
)
RE_PHONE_8 = re.compile(
    r'(?<!\d)(8[\s(-]*\d{3}[\s)-]*\d{3}[\s-]*\d{2}[\s-]*\d{2})(?!\d)'
)
# Локальные телефоны: (86559) 2-50-08, (495) 123-45-67
RE_PHONE_LOCAL = re.compile(
    r'(?:тел\.?|телефон|моб\.?|факс|т\.)\s*:?\s*'
    r'(\(\d{4,5}\)\s*\d[\d\s-]{4,12}\d)',
    re.IGNORECASE,
)
RE_PHONE_LOCAL_BARE = re.compile(
    r'(?<!\d)(\(\d{4,5}\)\s*\d[\d\s-]{4,12}\d)(?!\d)'
)

# Email
RE_EMAIL = re.compile(
    r'(?:e-?mail|эл\.?\s*почта|электронн\w*\s+почт\w*)\s*:?\s*'
    r'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)',
    re.IGNORECASE,
)
RE_EMAIL_BARE = re.compile(
    r'\b([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z]{2,})\b'
)

# Адрес — ловим максимально полно
RE_ADDRESS = re.compile(
    r'(?:ул(?:ица)?[-.\s]+|пр(?:оспект)?[-.\s]+|пер(?:еулок)?[-.\s]+|'
    r'б(?:ульвар)?[-.\s]+|(?:ш(?:оссе)?[-.\s]+)|наб(?:ережная)?[-.\s]+|'
    r'пл(?:ощадь)?[-.\s]+|мкр[-.\s]*н?[-.\s]+|проезд[-.\s]+)'
    r'[А-ЯЁа-яё][А-ЯЁа-яё0-9\s./-]{2,40}?'
    r'(?:,?\s*д(?:ом)?[-.\s]*\d+[а-яА-Я/]*)'
    r'(?:,?\s*(?:корп|стр|к)[-.\s]*\d+[а-яА-Я]*)?'
    r'(?:,?\s*(?:кв|оф|офис|пом|каб|комн|этаж)[-.\s]*\d+)?',
    re.IGNORECASE,
)

# Адрес без "д." — ул. Розы Люксембург, 1
RE_ADDRESS_SHORT = re.compile(
    r'(?:ул(?:ица)?[-.\s]+|пр(?:оспект)?[-.\s]+|пер(?:еулок)?[-.\s]+|'
    r'б(?:ульвар)?[-.\s]+|(?:ш(?:оссе)?[-.\s]+)|наб(?:ережная)?[-.\s]+|'
    r'пл(?:ощадь)?[-.\s]+|мкр[-.\s]*н?[-.\s]+|проезд[-.\s]+)'
    r'[А-ЯЁа-яё][А-ЯЁа-яё0-9\s./-]{2,40}?'
    r',\s*\d+[а-яА-Я/]*'
    r'(?:,?\s*(?:корп|стр|к)[-.\s]*\d+[а-яА-Я]*)?'
    r'(?:,?\s*(?:кв|оф|офис|пом|каб|комн|этаж)[-.\s]*\d+)?',
    re.IGNORECASE,
)

# Полный адрес с индексом: 356808, Россия, Ставропольский край, г. Буденновск, ул. Розы Люксембург, 1
RE_ADDRESS_FULL = re.compile(
    r'(?<!\d)\d{6}\s*,?\s*'                         # индекс
    r'(?:Росси[яи],?\s*)?'                          # Россия (опц.)
    r'(?:[А-ЯЁа-яё]+\s+(?:край|обл(?:асть)?|респ(?:ублика)?|округ|АО),?\s*)?'  # регион
    r'(?:(?:г\.|город|гор\.)\s*[А-ЯЁа-яё][-А-ЯЁа-яё\s]+,?\s*)?'  # город
    r'(?:(?:ул(?:ица)?|пр|пер|б-р|ш|наб|пл|проезд|мкр)[-.\s]+[А-ЯЁа-яё0-9][-А-ЯЁа-яё0-9\s./]+?'
    r'(?:,?\s*(?:д(?:ом)?[-.\s]*)?)\d+[а-яА-Я/]*)?'
    r'(?:,?\s*(?:корп|стр|к)[-.\s]*\d+)?'
    r'(?:,?\s*(?:кв|оф|офис|пом|каб|комн|этаж)[-.\s]*\d+)?',
    re.IGNORECASE,
)

# Юридический/фактический адрес — ловим всё до конца строки
RE_ADDRESS_LABEL = re.compile(
    r'(?:юридический|фактический|почтовый|адрес\s+местонахождения|'
    r'место\s+нахождения|адрес\s+регистрации)\s*(?:адрес)?\s*:?\s*'
    r'([^\n]{10,120})',
    re.IGNORECASE,
)

# Почтовый индекс
RE_INDEX = re.compile(r'(?<!\d)(\d{6})(?:,?\s+(?:г\.|город|Россия|РФ|обл\.|респ\.))', re.IGNORECASE)

# Город с префиксом
RE_CITY_PREFIX = re.compile(
    r'(?:г\.\s*|город\s+|гор\.\s*)([А-ЯЁ][а-яё]+(?:[-\s][А-ЯЁа-яё]+)?)',
    re.IGNORECASE,
)


# ── Автозамены (шаблоны заглушек) ───────────────────────────

_surname_counter = 0
_org_counter = 0
_city_counter = 0
_addr_counter = 0


def _reset_counters():
    global _surname_counter, _org_counter, _city_counter, _addr_counter
    _surname_counter = 0
    _org_counter = 0
    _city_counter = 0
    _addr_counter = 0


_replacement_cache: dict[str, str] = {}


def _auto_replacement(entity_type: str, original: str) -> str:
    """Генерирует автозаглушку для сущности."""
    global _surname_counter, _org_counter, _city_counter, _addr_counter

    key = f"{entity_type}:{original.strip().lower()}"
    if key in _replacement_cache:
        return _replacement_cache[key]

    if entity_type == ENTITY_SURNAME:
        _surname_counter += 1
        repl = f"Сотрудник №{_surname_counter}"
    elif entity_type == ENTITY_ORGANIZATION:
        _org_counter += 1
        repl = f'ООО «Организация-{_org_counter}»'
    elif entity_type == ENTITY_CITY:
        _city_counter += 1
        repl = f"г. Город-{_city_counter}"
    elif entity_type == ENTITY_INN:
        repl = "ИНН 0000000000"
    elif entity_type == ENTITY_OGRN:
        repl = "ОГРН 0000000000000"
    elif entity_type == ENTITY_KPP:
        repl = "КПП 000000000"
    elif entity_type == ENTITY_BIK:
        repl = "БИК 000000000"
    elif entity_type == ENTITY_ACCOUNT:
        repl = "00000000000000000000"
    elif entity_type == ENTITY_SNILS:
        repl = "000-000-000 00"
    elif entity_type == ENTITY_PASSPORT:
        repl = "серия 00 00 № 000000"
    elif entity_type == ENTITY_PHONE:
        repl = "+7 (000) 000-00-00"
    elif entity_type == ENTITY_EMAIL:
        repl = "email@example.com"
    elif entity_type == ENTITY_ADDRESS:
        _addr_counter += 1
        repl = f"ул. Улица, д. {_addr_counter}"
    else:
        repl = "[ЗАГЛУШКА]"

    _replacement_cache[key] = repl
    return repl


# ── Детекторы ────────────────────────────────────────────────

def _is_likely_surname(word: str) -> bool:
    """Проверяет, похоже ли слово на фамилию по морфологии."""
    if len(word) < 3:
        return False
    if not word[0].isupper():
        return False
    if not all(c.isalpha() for c in word):
        return False
    if word in FALSE_POSITIVE_SURNAMES:
        return False
    if word in RUSSIAN_CITIES:
        return False
    if word in ALL_FIRST_NAMES:
        return False

    lower = word.lower()

    # Фильтр прилагательных/причастий — типичные окончания
    _ADJ_ENDINGS = (
        'ский', 'ская', 'ское', 'ского', 'ской', 'ским', 'ском', 'ских',
        'ские', 'скую',
        'цкий', 'цкая', 'цкое', 'цкого', 'цкой', 'цким', 'цком', 'цких',
        'цкие', 'цкую',
        'ный', 'ная', 'ное', 'ного', 'ной', 'ным', 'ном', 'ных', 'ные', 'ную',
        'тый', 'тая', 'тое', 'того', 'той', 'тым', 'том', 'тых', 'тые', 'тую',
        'щий', 'щая', 'щее', 'щего', 'щей', 'щим', 'щем', 'щих', 'щие', 'щую',
        'чий', 'чья', 'чье', 'чьего', 'чьей', 'чьим', 'чьем', 'чьих', 'чьи',
        'жий', 'жая', 'жее', 'жего', 'жей', 'жим', 'жем', 'жих', 'жие',
        'ший', 'шая', 'шее', 'шего', 'шей', 'шим', 'шем', 'ших', 'шие',
        'ющих', 'ющий', 'ющая', 'ющее', 'ющего',
        'ящих', 'ящий', 'ящая', 'ящее', 'ящего',
        'вший', 'вшая', 'вшее', 'вшего', 'вших',
        'ический', 'ическая', 'ическое', 'ического', 'ических',
        'ильный', 'ильная', 'ильного',
        'альный', 'альная', 'альное', 'ального', 'альных',
        'ельный', 'ельная', 'ельное', 'ельного', 'ельных',
        'фиком', 'фика', 'фику', 'фике', 'фики',
    )
    for adj_end in _ADJ_ENDINGS:
        if lower.endswith(adj_end) and len(lower) > len(adj_end) + 1:
            return False

    # Фильтр существительных в родительном падеже множественного числа
    # "договоров", "объектов" и т.п. — не фамилии
    if lower.endswith('ов') and len(lower) > 5:
        # Проверяем что это не просто слово на -ов (типа "договоров")
        # Фамилии на -ов обычно короткие: Иванов (6), Петров (6)
        # Слова-не-фамилии длиннее: договоров (9), документов (10)
        if len(lower) > 8:
            return False

    for ending in ALL_SURNAME_ENDINGS:
        if lower.endswith(ending) and len(lower) > len(ending) + 1:
            return True
    return False


def detect_full_names(text: str) -> list[DetectedEntity]:
    """Обнаруживает полные ФИО: Фамилия Имя Отчество."""
    entities = []
    for m in RE_FULL_NAME.finditer(text):
        surname, name, patronymic = m.group(1), m.group(2), m.group(3)
        full = m.group(0)

        # Проверяем что имя — из справочника или похоже на имя
        name_ok = name in ALL_FIRST_NAMES or len(name) >= 3
        # Для полного ФИО (с отчеством) фамилия гарантированно верна —
        # отчество с -ович/-овна является надёжным маркером
        surname_ok = True  # отчество подтверждает что это ФИО

        if not name_ok:
            continue

        if is_whitelisted_org(full):
            continue

        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_SURNAME,
            replacement=_auto_replacement(ENTITY_SURNAME, surname),
            confidence=0.95,
        ))
    return entities


def detect_surname_initials(text: str) -> list[DetectedEntity]:
    """Обнаруживает Фамилия И.О. и И.О. Фамилия."""
    entities = []

    for m in RE_SURNAME_INITIALS.finditer(text):
        surname = m.group(1)
        if not _is_likely_surname(surname):
            continue
        full = m.group(0)
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_SURNAME,
            replacement=_auto_replacement(ENTITY_SURNAME, surname),
            confidence=0.95,
        ))

    for m in RE_INITIALS_SURNAME.finditer(text):
        surname = m.group(2)
        if not _is_likely_surname(surname):
            continue
        full = m.group(0)
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_SURNAME,
            replacement=_auto_replacement(ENTITY_SURNAME, surname),
            confidence=0.95,
        ))

    return entities


def detect_standalone_surnames(text: str, already_found: set[tuple] = None) -> list[DetectedEntity]:
    """
    Обнаруживает отдельные фамилии (без инициалов).
    Ищет слова с характерными окончаниями, проверяет контекст.
    already_found — множество (start, end) уже найденных сущностей.
    """
    if already_found is None:
        already_found = set()

    entities = []

    for m in re.finditer(r'\b([А-ЯЁ][а-яё]{2,})\b', text):
        word = m.group(1)
        start, end = m.start(), m.end()

        # Пропускаем уже найденные
        overlaps = False
        for fs, fe in already_found:
            if start >= fs and end <= fe:
                overlaps = True
                break
        if overlaps:
            continue

        if not _is_likely_surname(word):
            continue

        # Контекстная проверка: что вокруг?
        ctx_start = max(0, start - 50)
        ctx_end = min(len(text), end + 50)
        context = text[ctx_start:ctx_end]

        if is_whitelisted_in_context(word, context):
            continue

        # Пропускаем если слово — часть адреса (ул. Ленина, пр. Кирова)
        ctx_before = text[max(0, start - 20):start].lower()
        if any(mark in ctx_before for mark in (
            'ул.', 'ул ', 'улица', 'улицы', 'улице', 'улицу',
            'пр.', 'пр ', 'проспект', 'пер.', 'переулок',
            'б-р', 'бульвар', 'наб.', 'набережн', 'шоссе',
            'пл.', 'площад', 'им.', 'имени',
        )):
            continue

        # Повышаем уверенность если рядом есть признаки ФИО
        confidence = 0.7
        # Проверяем: стоит ли после слова имя из справочника?
        after = text[end:min(end + 30, len(text))]
        for name in ALL_FIRST_NAMES:
            if after.lstrip().startswith(name):
                confidence = 0.9
                break

        entities.append(DetectedEntity(
            start=start, end=end, text=word,
            entity_type=ENTITY_SURNAME,
            replacement=_auto_replacement(ENTITY_SURNAME, word),
            confidence=confidence,
        ))

    return entities


def detect_organizations(text: str) -> list[DetectedEntity]:
    """Обнаруживает организации: ООО «...», ИП ..., и т.д."""
    entities = []

    # ООО/ПАО/АО «Название»
    for pattern in (RE_ORG_QUOTED, RE_ORG_FULL_QUOTED):
        for m in pattern.finditer(text):
            full = m.group(0)
            if is_whitelisted_org(full):
                continue
            name_part = m.group(1).strip()
            if is_whitelisted_org(name_part):
                continue
            entities.append(DetectedEntity(
                start=m.start(), end=m.end(), text=full,
                entity_type=ENTITY_ORGANIZATION,
                replacement=_auto_replacement(ENTITY_ORGANIZATION, full),
                confidence=0.95,
            ))

    # ИП Фамилия
    for m in RE_IP.finditer(text):
        full = m.group(0)
        if is_whitelisted_org(full):
            continue
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_ORGANIZATION,
            replacement=_auto_replacement(ENTITY_ORGANIZATION, full),
            confidence=0.9,
        ))

    # Короткие названия: если нашли "ПАО «ЛУКОЙЛ»", ищем "ЛУКОЙЛ" отдельно
    short_names: dict[str, str] = {}  # short_name -> replacement
    for e in entities:
        # Извлекаем название из кавычек
        qm = re.search(r'[«""\']([^»""\'"]+)[»""\']', e.text)
        if qm:
            name = qm.group(1).strip()
            if len(name) >= 3 and name not in short_names:
                short_names[name] = e.replacement

    found_ranges = {(e.start, e.end) for e in entities}
    for name, repl in short_names.items():
        for m in re.finditer(r'\b' + re.escape(name) + r'\b', text):
            # Пропускаем если уже покрыто другой сущностью
            overlaps = False
            for (s, e) in found_ranges:
                if m.start() < e and m.end() > s:
                    overlaps = True
                    break
            if overlaps:
                continue
            if is_whitelisted_org(name):
                continue
            entities.append(DetectedEntity(
                start=m.start(), end=m.end(), text=name,
                entity_type=ENTITY_ORGANIZATION,
                replacement=repl,
                confidence=0.85,
            ))
            found_ranges.add((m.start(), m.end()))

    return entities


def detect_cities(text: str) -> list[DetectedEntity]:
    """Обнаруживает названия городов из справочника."""
    entities = []

    # г. Москва, город Москва
    for m in RE_CITY_PREFIX.finditer(text):
        city_name = m.group(1).strip()
        full = m.group(0)
        if city_name in RUSSIAN_CITIES:
            entities.append(DetectedEntity(
                start=m.start(), end=m.end(), text=full,
                entity_type=ENTITY_CITY,
                replacement=_auto_replacement(ENTITY_CITY, city_name),
                confidence=0.95,
            ))

    # Города без префикса — только в контексте адреса/реквизитов
    for city in RUSSIAN_CITIES:
        for m in re.finditer(r'\b' + re.escape(city) + r'\b', text):
            # Проверяем контекст: рядом должен быть адрес или индекс
            ctx_start = max(0, m.start() - 30)
            ctx_end = min(len(text), m.end() + 30)
            context = text[ctx_start:ctx_end].lower()

            in_address_ctx = any(marker in context for marker in (
                'адрес', 'ул.', 'ул ', 'улица', 'пр.', 'проспект',
                'пер.', 'переулок', 'обл.', 'область', 'край',
                'респ', 'район', 'индекс', ',', 'г.', 'город',
            ))

            if in_address_ctx:
                # Проверяем что не уже найдено с префиксом
                already = False
                for e in entities:
                    if m.start() >= e.start and m.end() <= e.end:
                        already = True
                        break
                if not already:
                    entities.append(DetectedEntity(
                        start=m.start(), end=m.end(), text=city,
                        entity_type=ENTITY_CITY,
                        replacement=_auto_replacement(ENTITY_CITY, city),
                        confidence=0.8,
                    ))

    return entities


def detect_requisites(text: str) -> list[DetectedEntity]:
    """Обнаруживает реквизиты: ИНН, ОГРН, КПП, БИК, счета."""
    entities = []

    for m in RE_INN.finditer(text):
        full = m.group(0)
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_INN,
            replacement=_auto_replacement(ENTITY_INN, m.group(1)),
        ))

    for m in RE_OGRN.finditer(text):
        full = m.group(0)
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_OGRN,
            replacement=_auto_replacement(ENTITY_OGRN, m.group(1)),
        ))

    for m in RE_KPP.finditer(text):
        full = m.group(0)
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_KPP,
            replacement=_auto_replacement(ENTITY_KPP, m.group(1)),
        ))

    for m in RE_BIK.finditer(text):
        full = m.group(0)
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=full,
            entity_type=ENTITY_BIK,
            replacement=_auto_replacement(ENTITY_BIK, m.group(1)),
        ))

    for pattern in (RE_ACCOUNT, RE_ACCOUNT_BARE):
        for m in pattern.finditer(text):
            full = m.group(0)
            entities.append(DetectedEntity(
                start=m.start(), end=m.end(), text=full,
                entity_type=ENTITY_ACCOUNT,
                replacement=_auto_replacement(ENTITY_ACCOUNT, m.group(1)),
            ))

    return entities


def detect_personal_ids(text: str) -> list[DetectedEntity]:
    """Обнаруживает СНИЛС, паспортные данные."""
    entities = []

    for m in RE_SNILS.finditer(text):
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=m.group(0),
            entity_type=ENTITY_SNILS,
            replacement=_auto_replacement(ENTITY_SNILS, m.group(1)),
        ))

    for m in RE_PASSPORT.finditer(text):
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=m.group(0),
            entity_type=ENTITY_PASSPORT,
            replacement=_auto_replacement(ENTITY_PASSPORT, m.group(0)),
        ))

    for m in RE_PASSPORT_ISSUED.finditer(text):
        entities.append(DetectedEntity(
            start=m.start(), end=m.end(), text=m.group(0),
            entity_type=ENTITY_PASSPORT,
            replacement="выдан [ОРГАН] ",
        ))

    return entities


def detect_contacts(text: str) -> list[DetectedEntity]:
    """Обнаруживает телефоны и email."""
    entities = []

    for pattern in (RE_PHONE, RE_PHONE_BARE, RE_PHONE_8, RE_PHONE_LOCAL, RE_PHONE_LOCAL_BARE):
        for m in pattern.finditer(text):
            entities.append(DetectedEntity(
                start=m.start(), end=m.end(), text=m.group(0),
                entity_type=ENTITY_PHONE,
                replacement=_auto_replacement(ENTITY_PHONE, m.group(0)),
            ))

    for pattern in (RE_EMAIL, RE_EMAIL_BARE):
        for m in pattern.finditer(text):
            email = m.group(0) if pattern == RE_EMAIL_BARE else m.group(0)
            entities.append(DetectedEntity(
                start=m.start(), end=m.end(), text=email,
                entity_type=ENTITY_EMAIL,
                replacement=_auto_replacement(ENTITY_EMAIL, email),
            ))

    return entities


def detect_addresses(text: str) -> list[DetectedEntity]:
    """Обнаруживает почтовые адреса."""
    entities = []

    # Полные адреса с индексом (самые длинные — первыми)
    for m in RE_ADDRESS_FULL.finditer(text):
        addr = m.group(0).strip().rstrip(',')
        if len(addr) < 10:
            continue
        entities.append(DetectedEntity(
            start=m.start(), end=m.start() + len(addr), text=addr,
            entity_type=ENTITY_ADDRESS,
            replacement=_auto_replacement(ENTITY_ADDRESS, addr),
        ))

    # Адрес по метке "Юридический адрес: ..."
    for m in RE_ADDRESS_LABEL.finditer(text):
        addr = m.group(1).strip().rstrip(',').rstrip('.')
        if len(addr) < 10:
            continue
        entities.append(DetectedEntity(
            start=m.start(1), end=m.start(1) + len(addr), text=addr,
            entity_type=ENTITY_ADDRESS,
            replacement=_auto_replacement(ENTITY_ADDRESS, addr),
        ))

    # Обычные адреса (ул. X, д. Y)
    for m in RE_ADDRESS.finditer(text):
        addr = m.group(0).strip().rstrip(',')
        if len(addr) < 5:
            continue
        entities.append(DetectedEntity(
            start=m.start(), end=m.start() + len(addr), text=addr,
            entity_type=ENTITY_ADDRESS,
            replacement=_auto_replacement(ENTITY_ADDRESS, addr),
        ))

    # Короткие адреса (ул. X, 1) — без "д."
    for m in RE_ADDRESS_SHORT.finditer(text):
        addr = m.group(0).strip().rstrip(',')
        if len(addr) < 5:
            continue
        entities.append(DetectedEntity(
            start=m.start(), end=m.start() + len(addr), text=addr,
            entity_type=ENTITY_ADDRESS,
            replacement=_auto_replacement(ENTITY_ADDRESS, addr),
        ))

    # Почтовый индекс отдельно
    for m in RE_INDEX.finditer(text):
        entities.append(DetectedEntity(
            start=m.start(), end=m.start() + 6, text=m.group(1),
            entity_type=ENTITY_ADDRESS,
            replacement="000000",
        ))

    return entities


# ── Основной API ─────────────────────────────────────────────

def auto_detect_all(text: str) -> list[DetectedEntity]:
    """
    Запускает все детекторы и возвращает дедуплицированный список сущностей,
    отсортированный по позиции в тексте.
    """
    _reset_counters()
    _replacement_cache.clear()

    all_entities: list[DetectedEntity] = []

    # Высокоточные детекторы (порядок важен — длинные совпадения первыми)
    all_entities.extend(detect_full_names(text))
    all_entities.extend(detect_surname_initials(text))
    all_entities.extend(detect_organizations(text))
    all_entities.extend(detect_requisites(text))
    all_entities.extend(detect_personal_ids(text))
    all_entities.extend(detect_contacts(text))
    all_entities.extend(detect_addresses(text))
    all_entities.extend(detect_cities(text))

    # Отдельные фамилии — с учётом уже найденных позиций
    found_ranges = {(e.start, e.end) for e in all_entities}
    all_entities.extend(detect_standalone_surnames(text, found_ranges))

    # Дедупликация: убираем пересечения, приоритет длинным
    return _deduplicate(all_entities)


def auto_detect_in_file(filepath: str) -> dict:
    """
    Запускает автодетекцию для файла.

    Возвращает:
    {
        "filepath": str,
        "entities": list[DetectedEntity],
        "pages": {page_num: text},  — для PDF
        "text": str,  — полный текст (для DOCX)
        "by_type": {type: [entities]},
    }
    """
    import os
    ext = os.path.splitext(filepath)[1].lower()

    _reset_counters()
    _replacement_cache.clear()

    if ext == '.docx':
        return _detect_in_docx(filepath)
    elif ext == '.pdf':
        return _detect_in_pdf(filepath)
    else:
        return {
            "filepath": filepath,
            "entities": [],
            "pages": {},
            "text": "",
            "by_type": {},
            "error": f"Неподдерживаемый формат: {ext}",
        }


def _detect_in_docx(filepath: str) -> dict:
    """Автодетекция в DOCX-файле."""
    from docx import Document

    try:
        doc = Document(filepath)
    except Exception as e:
        return {
            "filepath": filepath,
            "entities": [],
            "pages": {},
            "text": "",
            "by_type": {},
            "error": str(e),
        }

    # Собираем весь текст
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paragraphs.append(p.text)

    full_text = '\n'.join(paragraphs)
    entities = auto_detect_all(full_text)

    by_type = {}
    for e in entities:
        by_type.setdefault(e.entity_type, []).append(e)

    return {
        "filepath": filepath,
        "entities": entities,
        "pages": {1: full_text},
        "text": full_text,
        "by_type": by_type,
        "error": None,
    }


def _ocr_page(page) -> str:
    """OCR страницы PDF через tesseract (для сканов)."""
    try:
        import pytesseract
        from PIL import Image
        import io

        # Рендерим страницу в изображение (300 DPI)
        pix = page.get_pixmap(dpi=300)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))

        # OCR с русским языком
        text = pytesseract.image_to_string(img, lang='rus+eng')
        return text
    except Exception:
        return ""


def _is_scanned_page(page) -> bool:
    """Определяет, является ли страница сканом (мало текста, есть изображения)."""
    text = page.get_text().strip()
    # Если текста мало (менее 50 символов) — скорее всего скан
    if len(text) < 50:
        return True
    return False


def _detect_in_pdf(filepath: str) -> dict:
    """Автодетекция в PDF-файле. Поддерживает текстовые PDF и сканы (OCR)."""
    import fitz

    try:
        doc = fitz.open(filepath)
    except Exception as e:
        return {
            "filepath": filepath,
            "entities": [],
            "pages": {},
            "text": "",
            "by_type": {},
            "error": str(e),
        }

    pages = {}
    all_entities = []
    full_text_parts = []
    offset = 0
    used_ocr = False

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text()

        # Если страница выглядит как скан — пробуем OCR
        if _is_scanned_page(page):
            ocr_text = _ocr_page(page)
            if ocr_text.strip():
                page_text = ocr_text
                used_ocr = True

        pages[page_num + 1] = page_text

        # Детектим на каждой странице отдельно
        page_entities = auto_detect_all(page_text)

        # Корректируем позиции с учётом общего смещения
        for e in page_entities:
            e.start += offset
            e.end += offset

        all_entities.extend(page_entities)
        full_text_parts.append(page_text)
        offset += len(page_text)

    doc.close()
    full_text = ''.join(full_text_parts)

    # Дедупликация глобальная
    all_entities = _deduplicate(all_entities)

    by_type = {}
    for e in all_entities:
        by_type.setdefault(e.entity_type, []).append(e)

    return {
        "filepath": filepath,
        "entities": all_entities,
        "pages": pages,
        "text": full_text,
        "by_type": by_type,
        "used_ocr": used_ocr,
        "error": None,
    }


def _deduplicate(entities: list[DetectedEntity]) -> list[DetectedEntity]:
    """Убирает пересекающиеся сущности, приоритет более длинным."""
    if not entities:
        return []

    # Сортируем: сначала по позиции, при равной позиции — длинные первыми
    entities.sort(key=lambda e: (e.start, -(e.end - e.start)))

    result = []
    last_end = -1
    for e in entities:
        if e.start >= last_end:
            result.append(e)
            last_end = e.end

    return result


# ── Утилита: типы → русские названия ────────────────────────

ENTITY_TYPE_NAMES = {
    ENTITY_SURNAME: "Фамилии / ФИО",
    ENTITY_ORGANIZATION: "Организации",
    ENTITY_CITY: "Города",
    ENTITY_INN: "ИНН",
    ENTITY_OGRN: "ОГРН",
    ENTITY_KPP: "КПП",
    ENTITY_BIK: "БИК",
    ENTITY_ACCOUNT: "Банковские счета",
    ENTITY_SNILS: "СНИЛС",
    ENTITY_PASSPORT: "Паспортные данные",
    ENTITY_PHONE: "Телефоны",
    ENTITY_EMAIL: "Email",
    ENTITY_ADDRESS: "Адреса",
}


def get_type_name(entity_type: str) -> str:
    """Возвращает русское название типа сущности."""
    return ENTITY_TYPE_NAMES.get(entity_type, entity_type)
